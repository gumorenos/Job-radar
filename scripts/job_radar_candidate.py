#!/usr/bin/env python3
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[1]
PROFILE_DIR = ROOT / "tracking" / "job-radar" / "profile"
CV_MARKDOWN_PATH = PROFILE_DIR / "cv.md"
CANDIDATE_PROFILE_PATH = PROFILE_DIR / "candidate-profile.json"

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}

ROLE_PATTERNS = [
    "HR Business Partner",
    "HRBP",
    "People Analytics",
    "People Operations",
    "Human Resources Manager",
    "Human Resources",
    "Recursos Humanos",
    "Gestión de Personas",
    "People Management",
    "HR Operations",
    "Compensation",
    "Benefits",
    "Total Rewards",
    "Talent Management",
    "Labor Relations",
    "Desarrollo Organizacional",
    "Organizational Development",
]

SKILL_PATTERNS = [
    "Power BI",
    "Excel",
    "SQL",
    "Python",
    "Analytics",
    "Data",
    "Dashboard",
    "KPIs",
    "OKR",
    "Compensation",
    "Benefits",
    "Recruiting",
    "Selection",
    "Employee Relations",
    "Labor Relations",
    "Labor Cost",
    "HR Budgeting",
    "Change Management",
    "Culture",
    "Talent",
    "Onboarding",
    "Payroll",
    "Job Evaluation",
    "Organizational Transformation",
    "HRIS",
    "Workday",
    "SAP",
    "SuccessFactors",
]

INDUSTRY_PATTERNS = [
    "retail",
    "industrial",
    "fintech",
    "technology",
    "consulting",
    "banking",
    "healthcare",
    "mining",
    "consumer goods",
    "manufacturing",
]

LANGUAGE_PATTERNS = ["English", "Inglés", "Spanish", "Español", "Portuguese", "Portugués"]

NEGATIVE_DEFAULTS = ["junior", "intern", "practicante", "assistant", "asistente", "sales"]


def clean_text(value: Any, limit: int = 20000) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).strip())[:limit]


def unique(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        normalized = clean_text(item, 120)
        key = normalized.lower()
        if normalized and key not in seen:
            seen.add(key)
            result.append(normalized)
    return result


def markdown_from_pdf(path: Path) -> str:
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("Falta PyMuPDF para procesar PDF. Instala `pymupdf`.") from exc
    doc = fitz.open(path)
    sections = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if text:
            sections.append(f"## Página {page_num}\n\n{text}")
    return "\n\n".join(sections).strip()


def markdown_from_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("Falta python-docx para procesar DOCX. Instala `python-docx`.") from exc
    doc = Document(path)
    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if "heading" in style or "title" in style:
            lines.append(f"## {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n\n".join(lines).strip()


def convert_to_markdown(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return markdown_from_pdf(path)
    if ext == ".docx":
        return markdown_from_docx(path)
    if ext in {".md", ".markdown", ".txt"}:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    raise ValueError(f"Formato no soportado: {ext}")


def pattern_hits(markdown: str, patterns: list[str]) -> list[str]:
    text = markdown.lower()
    hits = []
    for pattern in patterns:
        needle = pattern.lower()
        if re.search(rf"(?<![a-z0-9]){re.escape(needle)}(?![a-z0-9])", text):
            hits.append(pattern)
    return unique(hits)


def infer_years(markdown: str) -> int | None:
    matches = re.findall(r"(\d{1,2})\+?\s*(?:años|years)", markdown, re.I)
    if not matches:
        return None
    values = [int(match) for match in matches if int(match) < 50]
    return max(values) if values else None


def extract_candidate_profile(markdown: str, source_filename: str = "") -> dict:
    roles = pattern_hits(markdown, ROLE_PATTERNS)
    skills = pattern_hits(markdown, SKILL_PATTERNS)
    industries = pattern_hits(markdown, INDUSTRY_PATTERNS)
    languages = pattern_hits(markdown, LANGUAGE_PATTERNS)
    years = infer_years(markdown)
    role_terms = unique([role.lower() for role in roles] + [skill.lower() for skill in skills[:8]])
    return {
        "source_filename": source_filename,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "target_roles": roles,
        "role_terms": role_terms,
        "skills": skills,
        "industries": industries,
        "languages": languages,
        "years_experience": years,
        "negative_terms": NEGATIVE_DEFAULTS,
    }


def load_candidate_profile() -> dict:
    if not CANDIDATE_PROFILE_PATH.exists():
        return {}
    return json.loads(CANDIDATE_PROFILE_PATH.read_text(encoding="utf-8"))


def save_candidate_profile(markdown: str, profile: dict) -> dict:
    PROFILE_DIR.mkdir(parents=True, exist_ok=True)
    CV_MARKDOWN_PATH.write_text(markdown.strip() + "\n", encoding="utf-8")
    CANDIDATE_PROFILE_PATH.write_text(json.dumps(profile, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return profile


def process_cv_upload(filename: str, content: bytes) -> dict:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Formato no soportado: {ext}. Usa PDF, DOCX, MD o TXT.")
    PROFILE_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(filename).name)[:160] or f"cv{ext}"
    original_path = PROFILE_DIR / f"original-{safe_name}"
    original_path.write_bytes(content)
    markdown = convert_to_markdown(original_path)
    profile = extract_candidate_profile(markdown, safe_name)
    save_candidate_profile(markdown, profile)
    return {
        "ok": True,
        "original_path": str(original_path),
        "markdown_path": str(CV_MARKDOWN_PATH),
        "profile_path": str(CANDIDATE_PROFILE_PATH),
        "markdown": markdown,
        "candidate_profile": profile,
    }
